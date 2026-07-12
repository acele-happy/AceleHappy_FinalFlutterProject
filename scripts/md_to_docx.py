#!/usr/bin/env python3
"""Convert the technical report markdown to .docx."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "StudentName_FinalFlutterProject.md"
DOCX_PATH = ROOT / "StudentName_FinalFlutterProject.docx"


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_formatted_run(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic


def parse_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            add_formatted_run(paragraph, token[2:-2], bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif token.startswith("["):
            label = re.match(r"\[([^\]]+)\]", token).group(1)
            paragraph.add_run(label)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            clean = cell_text.strip()
            if clean.startswith("**") and clean.endswith("**"):
                add_formatted_run(p, clean[2:-2], bold=True)
            else:
                parse_inline(p, clean)


def convert(md_text: str, output_path: Path) -> None:
    doc = Document()
    set_default_style(doc)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                in_code = False
                code_lines = []
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            in_code = True
            code_lines = []
            i += 1
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(set(c) <= {"-", ":"} for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                add_table(doc, table_rows)
                table_rows = []
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, line[6:].strip())
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, line[2:].strip())
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, re.sub(r"^\d+\.\s", "", line).strip())
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(line[2:].strip())
            run.italic = True
        else:
            p = doc.add_paragraph()
            parse_inline(p, line.strip())

        i += 1

    doc.save(output_path)


def main() -> int:
    md_path = Path(sys.argv[1]) if len(sys.argv) > 1 else MD_PATH
    docx_path = Path(sys.argv[2]) if len(sys.argv) > 2 else DOCX_PATH

    if not md_path.exists() or md_path.stat().st_size == 0:
        print(f"Error: markdown file missing or empty: {md_path}", file=sys.stderr)
        return 1

    convert(md_path.read_text(encoding="utf-8"), docx_path)
    print(f"Created: {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
